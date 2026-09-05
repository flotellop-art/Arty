"""Independent reader of synthetic W05 fixtures; never opens user documents.

Generate with ARTY_OFFICE_FIXTURE_DIR and officeExport.test.ts, then pass that
directory here. Structural validation only, not Word/Excel visual rendering.
"""
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET
from docx import Document
from openpyxl import load_workbook
import docx
import openpyxl

directory = Path(sys.argv[1]).resolve()
word = Document(directory / "arty-editable.docx")
text = "\n".join(p.text for p in word.paragraphs)
assert "Été 😀" in text
assert "Réponse interrompue" in text
assert "document-1.txt" in text and "document-2.txt" in text
assert len(word.tables) == 2
assert round(word.sections[0].page_width.twips) == 11906
assert round(word.sections[0].page_height.twips) == 16838
assert word.tables[0].cell(3, 0).text == "0012"

book = load_workbook(directory / "arty-editable.xlsx", data_only=False)
assert book.sheetnames == ["Tableau 1 - M1", "Tableau 2 - M2", "Informations"]
sheet = book.worksheets[0]
expected = ['=HYPERLINK("https://example.test")', '+33123456789', '0012',
            '1234567890123456', '_x0041_', '1,234', '1.234', '1e3', '@SUM(A1)']
for row, value in enumerate(expected, 2):
    cell = sheet.cell(row, 1)
    # openpyxl 3.1.5 preserves ST_Xstring encoding in inlineStr; Excel decodes it.
    expected_value = '_x005F_x0041_' if value == '_x0041_' else value
    assert cell.value == expected_value, (row, cell.value, expected_value)
    assert cell.data_type == 's', (row, cell.data_type)
    assert cell.number_format == '@'
for worksheet in book:
    for row in worksheet:
        assert all(cell.data_type != 'f' for cell in row)

for name in ['arty-editable.docx', 'arty-editable.xlsx']:
    with ZipFile(directory / name) as archive:
        for path in archive.namelist():
            assert not any(part in path.lower() for part in ['vbaproject', 'externallinks', 'embeddings'])
            if path.endswith(('.xml', '.rels')):
                root = ET.fromstring(archive.read(path))
                for node in root.iter():
                    assert node.tag.split('}')[-1] not in {'f', 'definedNames', 'altChunk', 'oleObject'}
                    assert node.attrib.get('TargetMode') != 'External'
print(f'PASS independent structure/content: python-docx {docx.__version__}, openpyxl {openpyxl.__version__}; no visual claim.')
