"""Emit synthetic OOXML fixtures from independent producers, never user data.

Run with python-docx and openpyxl installed. Output JSON may be committed as
src/__tests__/helpers/office-producer-fixtures.json. No files are read/written.
"""
import base64
import io
import json
from datetime import datetime

import docx
import openpyxl

document = docx.Document()
document.core_properties.author = "Arty synthetic test"
document.core_properties.created = datetime(2026, 9, 5)
document.core_properties.modified = datetime(2026, 9, 5)
document.add_heading("Projet été", level=1)
document.add_paragraph("Facture synthétique : 1250 € — 東京 🌞")
table = document.add_table(rows=2, cols=2)
table.cell(0, 0).text = "Client"
table.cell(0, 1).text = "Montant"
table.cell(1, 0).text = "Exemple"
table.cell(1, 1).text = "1250"
word = io.BytesIO()
document.save(word)

workbook = openpyxl.Workbook()
workbook.properties.creator = "Arty synthetic test"
workbook.properties.created = datetime(2026, 9, 5)
workbook.properties.modified = datetime(2026, 9, 5)
sheet = workbook.active
sheet.title = "Ventes été"
sheet.append(["Client", "Montant"])
sheet.append(["Exemple 東京", 1250])
sheet["B3"] = "=SUM(B2:B2)"
hidden = workbook.create_sheet("Notes")
hidden.sheet_state = "hidden"
hidden["A1"] = "Brouillon synthétique"
excel = io.BytesIO()
workbook.save(excel)

print(json.dumps({
    "provenance": {"synthetic": True, "python_docx": docx.__version__, "openpyxl": openpyxl.__version__},
    "docx": base64.b64encode(word.getvalue()).decode("ascii"),
    "xlsx": base64.b64encode(excel.getvalue()).decode("ascii"),
}))
